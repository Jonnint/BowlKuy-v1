"""
Script untuk convert PROPOSAL.md ke format Word (.docx)
Install dulu: pip install python-docx markdown
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_word_document():
    # Baca file markdown
    with open('PROPOSAL.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Buat dokumen Word
    doc = Document()
    
    # Set margin
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Split content by lines
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            continue
        
        # Heading 1 (# )
        if line.startswith('# ') and not line.startswith('## '):
            text = line.replace('# ', '')
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Heading 2 (## )
        elif line.startswith('## ') and not line.startswith('### '):
            text = line.replace('## ', '')
            doc.add_heading(text, level=2)
            
        # Heading 3 (### )
        elif line.startswith('### '):
            text = line.replace('### ', '')
            doc.add_heading(text, level=3)
            
        # Horizontal rule
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)
            
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = re.sub(r'^[-*] ', '', line)
            # Remove markdown bold
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Bullet')
            
        # Numbered list
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Number')
            
        # Table detection (simple)
        elif line.startswith('|'):
            # Skip table for now, add as paragraph
            text = line.replace('|', ' | ')
            p = doc.add_paragraph(text)
            run = p.runs[0]
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            
        # Code block
        elif line.startswith('```'):
            continue
            
        # Regular paragraph
        else:
            # Remove markdown bold
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            if text:
                doc.add_paragraph(text)
    
    # Save document
    doc.save('PROPOSAL_UPDATED.docx')
    print("✓ File PROPOSAL_UPDATED.docx berhasil dibuat!")
    print("✓ Lokasi: PROPOSAL_UPDATED.docx")

if __name__ == '__main__':
    try:
        create_word_document()
    except Exception as e:
        print(f"Error: {e}")
        print("\nPastikan sudah install: pip install python-docx")
